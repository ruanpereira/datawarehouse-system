from typing import Any, Dict
from docx import Document
from docx.shared import Pt

class ReportGenerator:
    def __init__(self, styles: Dict[str, Any] = None) -> None:
        self.styles = styles or {'font_name': 'Arial','font_size': Pt(11)}

    def _config_styles(self, doc: Document) -> None:
        style = doc.styles['Normal']
        font = style.font  # type: ignore
        font.name = self.styles['font_name']
        font.size = self.styles['font_size']

    def generate(self, data: Dict[str, Any], filename: str) -> None:
        doc = Document()
        self._config_styles(doc)
        self._add_heading(doc, 'Relatório de Vendas', level=1)
        filtro = data.get('filtro')
        if filtro:
            self._add_paragraph(doc, f"Filtro aplicado: {filtro}")
        self._add_paragraph(doc, f"Vendedor: {data.get('nome_vendedor','N/A')}")
        self._add_section(doc, 'Resumo Geral', data.get('media_total'))
        for key, title in [('media_regiao','Média por Região'),('media_periodo','Média por Período'),('media_data','Média por Data')]:
            if data.get(key): self._add_table(doc,title,data[key])
        if data.get('media_geral'): self._add_section(doc,'Média Geral de Todos os Vendedores',data['media_geral'])
        self._add_section(doc,'Observações Finais',data.get('observacoes',''))
        doc.save(filename)

    def _add_heading(self, doc: Document, text: str, level: int) -> None:
        run = doc.add_heading(level=level).add_run(text)
        run.bold = True

    def _add_paragraph(self, doc: Document, text: str) -> None:
        doc.add_paragraph(text)

    def _add_section(self, doc: Document, title: str, content: Any) -> None:
        self._add_heading(doc, title, level=2)
        self._add_paragraph(doc, str(content))

    def _add_table(self, doc: Document, title: str, data: Dict[Any, Any]) -> None:
        self._add_heading(doc, title, level=2)
        table = doc.add_table(rows=1, cols=2)
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text = 'Categoria','Valor Médio'
        for cat, val in data.items():
            cells = table.add_row().cells
            cells[0].text = str(cat)
            cells[1].text = str(val)