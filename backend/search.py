import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import pandas as pd
import os
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Pt

# Classes originais adaptadas (importadas)

class RelatorioVendas:
    def __init__(self, dados):
        self.doc = Document()
        self.dados = dados
        self._configurar_estilos()

    def _configurar_estilos(self):
        style = self.doc.styles['Normal']
        font = style.font  # type: ignore
        font.name = 'Arial'
        font.size = Pt(11)

    def add_title(self, text):
        p = self.doc.add_heading(level=1)
        run = p.add_run(text)
        run.bold = True

    def add_subtitle(self, text):
        p = self.doc.add_heading(level=2)
        run = p.add_run(text)
        run.bold = True

    def criar_relatorio(self):
        self.add_title("Relatório de Vendas")
        self.doc.add_paragraph(f"Vendedor: {self.dados['nome_vendedor']}")
        self.add_subtitle("Resumo Geral")
        self.doc.add_paragraph(f"Média total de vendas: {self.dados['media_total']}")
        self._criar_tabela("Média por Região", 'media_regiao')
        self._criar_tabela("Média por Período", 'media_periodo')
        if self.dados.get('media_data'):
            self._criar_tabela("Média por Data", 'media_data')
        if self.dados.get('media_geral'):
            self.add_subtitle("Média Geral de Todos os Vendedores")
            self.doc.add_paragraph(f"Média geral: {self.dados['media_geral']}")
        self.add_subtitle("Observações Finais")
        self.doc.add_paragraph("[Análises adicionais e comentários]")
        self.doc.save('relatorio_vendas_completo.docx')

    def _criar_tabela(self, titulo, chave):
        self.add_subtitle(titulo)
        table = self.doc.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Categoria'
        hdr_cells[1].text = 'Valor Médio'
        for chave_item, valor in self.dados[chave].items():
            row_cells = table.add_row().cells
            row_cells[0].text = chave_item
            row_cells[1].text = valor
class SearchingFile:
    def __init__(self):
        self.dataset = None
        self.columns_list = []
        self.numeric_columns = []
        self.relatorio_data = {}

    def load_dataset(self, file_path):
        """Carrega diferentes formatos de arquivo"""
        ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if ext == '.csv':
                self.dataset = pd.read_csv(file_path)
            elif ext in ('.xlsx', '.xls'):
                self.dataset = pd.read_excel(file_path)
            elif ext == '.pdf':
                reader = PdfReader(file_path)
                text = "\n".join([page.extract_text() for page in reader.pages])
                self.dataset = pd.DataFrame({'Texto': [text]})
            else:
                raise ValueError("Formato não suportado")
            
            self._analisar_colunas()
            return True
        except Exception as e:
            messagebox.showerror("Erro", f"Falha ao ler arquivo: {str(e)}")
            return False

    def _analisar_colunas(self):
        """Identifica colunas e tipos disponíveis"""
        self.columns_list = self.dataset.columns.tolist()
        self.numeric_columns = self.dataset.select_dtypes(include=['number']).columns.tolist()
        
        # Se não encontrar colunas numéricas, tenta converter algumas comuns
        if not self.numeric_columns:
            for col in ['valor', 'venda', 'preço', 'total']:
                if col in [c.lower() for c in self.columns_list]:
                    col_real = [c for c in self.columns_list if c.lower() == col][0]
                    try:
                        self.dataset[col_real] = pd.to_numeric(self.dataset[col_real])
                        self.numeric_columns.append(col_real)
                    except:
                        continue

    def preparar_relatorio(self, coluna, valor, coluna_media=None, agrupamentos=None):
        if 'Texto' in self.columns_list:
            messagebox.showwarning("Aviso", "PDF não suporta análise detalhada. Gerando relatório básico.")
            self.relatorio_data = {'conteudo_pdf': self.dataset['Texto'].iloc[0]}
            self._gerar_relatorio_pdf()
            return

        # Configura valores padrão se não fornecidos
        coluna_media = coluna_media or (self.numeric_columns[0] if self.numeric_columns else None)
        agrupamentos = agrupamentos or [c for c in self.columns_list if c != coluna]
        
        if not coluna_media:
            messagebox.showerror("Erro", "Nenhuma coluna numérica encontrada para cálculo de médias")
            return

        # Filtra os dados
        self.column_name = coluna
        self.selected_value = valor
        self.filtered_rows = self.dataset[self.dataset[self.column_name] == self.selected_value]
        
        # Prepara dados do relatório
        self.relatorio_data = {
            'filtro': f"{coluna}: {valor}",
            'dados_gerais': self._get_dados_gerais(coluna_media),
            'agrupamentos': {}
        }

        # Calcula médias para cada agrupamento
        for agrupamento in agrupamentos:
            if agrupamento in self.columns_list and agrupamento != coluna:
                self.relatorio_data['agrupamentos'][agrupamento] = self._formatar_medias(agrupamento, coluna_media)
        
        self._gerar_relatorio_docx()

    def _get_dados_gerais(self, coluna_media):
        """Calcula estatísticas básicas"""
        return {
            'média': f"R$ {self.filtered_rows[coluna_media].mean():.2f}",
            'média_geral': f"R$ {self.dataset[coluna_media].mean():.2f}",
            'total': f"R$ {self.filtered_rows[coluna_media].sum():.2f}",
            'contagem': len(self.filtered_rows)
        }

    def _formatar_medias(self, agrupamento, coluna_media):
        """Formata médias para um agrupamento específico"""
        try:
            return self.filtered_rows.groupby(agrupamento)[coluna_media].mean().apply(
                lambda x: f"R$ {x:.2f}" if pd.notnull(x) else "N/A"
            ).to_dict()
        except:
            return {"Erro": "Não foi possível calcular agrupamento"}

    def _gerar_relatorio_docx(self):
        """Gera relatório DOCX dinâmico"""
        doc = Document()
        
        # Cabeçalho
        doc.add_heading('Relatório Analítico', 0)
        doc.add_paragraph(f"Filtro aplicado: {self.relatorio_data['filtro']}")
        
        # Dados gerais
        doc.add_heading('Métricas Gerais', level=1)
        for k, v in self.relatorio_data['dados_gerais'].items():
            doc.add_paragraph(f"{k.replace('_', ' ').title()}: {v}")
        
        # Agrupamentos
        for agrupamento, dados in self.relatorio_data['agrupamentos'].items():
            doc.add_heading(f"Por {agrupamento}", level=2)
            table = doc.add_table(rows=1, cols=2)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = agrupamento
            hdr_cells[1].text = 'Média'
            
            for categoria, media in dados.items():
                row_cells = table.add_row().cells
                row_cells[0].text = str(categoria)
                row_cells[1].text = media
        
        doc.save('relatorio_analitico.docx')
        messagebox.showinfo("Sucesso", "Relatório gerado com sucesso!")

    def _gerar_relatorio_pdf(self):
        """Gera relatório básico para PDF"""
        doc = Document()
        doc.add_heading('Relatório Básico PDF', 0)
        doc.add_paragraph(self.relatorio_data['conteudo_pdf'][:5000]) 
        messagebox.showinfo("Sucesso", "Relatório básico do PDF gerado!")
# Interface Gráfica

class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Gerador de Relatório de Vendas")
        self.geometry("600x450")

        self.searching = SearchingFile()

        self._criar_widgets()

    def _criar_widgets(self):
        frame = ttk.Frame(self, padding=10)
        frame.pack(expand=True, fill='both')

        ttk.Label(frame, text="Selecione o arquivo de dados:").grid(row=0, column=0, pady=5, sticky='w')
        self.botao_arquivo = ttk.Button(frame, text="Procurar Arquivo", command=self.carregar_arquivo)
        self.botao_arquivo.grid(row=1, column=0, pady=5, sticky='ew')

        ttk.Label(frame, text="Coluna para filtrar:").grid(row=2, column=0, pady=5, sticky='w')
        self.combo_colunas = ttk.Combobox(frame, state="readonly")
        self.combo_colunas.grid(row=3, column=0, pady=5, sticky='ew')
        self.combo_colunas.bind("<<ComboboxSelected>>", self.carregar_valores)

        ttk.Label(frame, text="Valor para filtrar:").grid(row=4, column=0, pady=5, sticky='w')
        self.combo_valores = ttk.Combobox(frame, state="readonly")
        self.combo_valores.grid(row=5, column=0, pady=5, sticky='ew')

        self.chk_media_data = tk.BooleanVar()
        self.chk_media_geral = tk.BooleanVar()
        ttk.Checkbutton(frame, text="Incluir média por data", variable=self.chk_media_data).grid(row=6, column=0, pady=5, sticky='w')
        ttk.Checkbutton(frame, text="Incluir média geral", variable=self.chk_media_geral).grid(row=7, column=0, pady=5, sticky='w')

        ttk.Button(frame, text="Gerar Relatório", command=self.gerar_relatorio).grid(row=8, column=0, pady=20, sticky='ew')

    def carregar_arquivo(self):
        tipos_arquivo = [
            ("Arquivos Suportados", "*.csv *.xlsx *.xls *.pdf"),
            ("CSV", "*.csv"), 
            ("Excel", "*.xlsx *.xls"),
            ("PDF", "*.pdf")
        ]
        
        file_path = filedialog.askopenfilename(filetypes=tipos_arquivo)
        if file_path and self.searching.load_dataset(file_path):
            self.combo_colunas['values'] = self.searching.columns_list
            self._atualizar_status(f"Arquivo carregado: {os.path.basename(file_path)}")

    def carregar_valores(self, event):
        coluna = self.combo_colunas.get()
        if coluna and hasattr(self.searching, 'dataset'):
            valores = sorted(self.searching.dataset[coluna].dropna().unique().tolist())
            self.combo_valores['values'] = valores
            self._atualizar_status(f"{len(valores)} valores encontrados na coluna {coluna}")

    def gerar_relatorio(self):
        if not self._validar_selecoes():
            return
        
        try:
            self.searching.preparar_relatorio(
                self.combo_colunas.get(),
                self.combo_valores.get(),
                self.chk_media_data.get(),
                self.chk_media_geral.get()
            )
        except Exception as e:
            messagebox.showerror("Erro", f"Falha ao gerar relatório: {str(e)}")

    def _validar_selecoes(self):
        if not (self.combo_colunas.get() and self.combo_valores.get()):
            messagebox.showerror("Erro", "Selecione a coluna e o valor para filtrar!")
            return False
        return True

    def _atualizar_status(self, mensagem):
        self.botao_arquivo['text'] = mensagem[:30] + "..." if len(mensagem) > 30 else mensagem

if __name__ == "__main__":
    app = App()
    app.mainloop()