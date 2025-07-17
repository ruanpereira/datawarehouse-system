from typing import Any, Dict, List, Optional
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import pandas as pd

class ReportGenerator:
    """
    Classe para gerar relatórios em formato Word com base em dados estruturados.
    
    Atributos:
        styles (Dict[str, Any]): Configurações de estilo para o documento
        _SECTION_CONFIGS (List[Tuple]): Configurações para seções dinâmicas
    """
    _SECTION_CONFIGS = [
        ('media_regiao', 'Média por Região'),
        ('media_periodo', 'Média por Período'),
        ('media_data', 'Média por Data')
    ]

    def __init__(self, styles: Optional[Dict[str, Any]] = None) -> None:
        """
        Inicializa o gerador de relatórios com estilos opcionais.
        
        Args:
            styles (Dict[str, Any], opcional): Dicionário com configurações de estilo
        """
        self.styles = styles or {
            'font_name': 'Arial',
            'font_size': Pt(11),
            'heading_margins': Inches(0.5)
        }

    def _config_styles(self, doc: Document) -> None:
        """Configura os estilos padrão do documento."""
        style = doc.styles['Normal']
        font = style.font
        font.name = self.styles['font_name']
        font.size = self.styles['font_size']

    def _add_heading(self, doc: Document, text: str, level: int) -> None:
        """
        Adiciona um cabeçalho ao documento com estilo em negrito.
        
        Args:
            doc (Document): Documento do Word
            text (str): Texto do cabeçalho
            level (int): Nível do cabeçalho (1-6)
        """
        heading = doc.add_heading(level=level)
        run = heading.add_run(text)
        run.bold = True
        self._set_heading_margins(heading)

    def _set_heading_margins(self, heading) -> None:
        """Define margens consistentes para os cabeçalhos."""
        prs = heading.paragraph_format
        prs.space_before = self.styles['heading_margins']
        prs.space_after = self.styles['heading_margins']

    def _add_paragraph(self, doc: Document, text: str) -> None:
        """Adiciona um parágrafo de texto simples."""
        doc.add_paragraph(text)

    def _add_section(self, doc: Document, title: str, content: Any) -> None:
        """
        Adiciona uma seção ao documento com título e conteúdo.
        
        Args:
            doc (Document): Documento do Word
            title (str): Título da seção
            content (Any): Conteúdo da seção (será convertido para string)
        """
        if content is None:
            return
        self._add_heading(doc, title, level=2)
        self._add_paragraph(doc, str(content))

    def _add_table(self, doc: Document, title: str, data: Dict[str, float], 
                  headers: List[str] = ['Categoria', 'Valor Médio']) -> None:
        """
        Adiciona uma tabela formatada ao documento.
        
        Args:
            doc (Document): Documento do Word
            title (str): Título da tabela
            data (Dict): Dados para a tabela
            headers (List[str]): Cabeçalhos das colunas
        """
        self._add_heading(doc, title, level=2)
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Light Shading Accent 1'
        
        # Configurar cabeçalho
        hdr_row = table.rows[0]
        for i, header in enumerate(headers):
            cell = hdr_row.cells[i]
            cell.text = header
            self._set_cell_bold(cell)

        # Adicionar dados
        for category, value in data.items():
            row_cells = table.add_row().cells
            row_cells[0].text = str(category)
            row_cells[1].text = f"{value:.2f}" if isinstance(value, float) else str(value)

        # Ajustar largura das colunas
        for column in table.columns:
            column.width = Inches(2.5)

    def _set_cell_bold(self, cell) -> None:
        """Define o texto da célula em negrito."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tag = qn('w:b')
        element = OxmlElement(tag)
        element.set(qn('w:val'), 'true')
        tcPr.append(element)

    def generate(self, data: Dict[str, Any], filename: str) -> None:
        """
        Gera o documento completo com base nos dados fornecidos.
        
        Args:
            data (Dict): Dados estruturados para o relatório
            filename (str): Nome do arquivo de saída
        """
        doc = Document()
        self._config_styles(doc)
        
        # Cabeçalho principal
        self._add_heading(doc, 'Relatório de Vendas', level=1)
        
        # Informações básicas
        if filtro := data.get('filtro'):
            self._add_paragraph(doc, f"Filtro aplicado: {filtro}")
            
        if vendedor := data.get('nome_vendedor'):
            self._add_paragraph(doc, f"Vendedor: {vendedor}")
        
        # Seções dinâmicas
        if media_total := data.get('media_total'):
            self._add_section(doc, 'Resumo Geral', media_total)
            
        for key, title in self._SECTION_CONFIGS:
            if table_data := data.get(key):
                self._add_table(doc, title, table_data)
        
        if media_geral := data.get('media_geral'):
            self._add_section(doc, 'Média Geral de Todos os Vendedores', media_geral)
            
        self._add_section(doc, 'Observações Finais', data.get('observacoes'))
        
        doc.save(filename)

class RelatorioToDf:
    """
    Classe para converter dados de relatório em DataFrame.
    
    Métodos:
        generate: Converte dados de relatório estruturados em DataFrame
    """
    @staticmethod
    def generate(relatorio: Dict[str, Dict[str, Any]]) -> pd.DataFrame:
        """
        Converte dados de relatório em DataFrame consolidado.

        Args:
        relatorio (Dict): Dados estruturados do relatório

        Returns:
        pd.DataFrame: DataFrame consolidado
        """
        dfs = []
        for consorciado, data in relatorio.items():
            vendedores = data['vendedores']

            if isinstance(vendedores, dict):
                if all(isinstance(v, list) for v in vendedores.values()):
                    df = pd.DataFrame(vendedores)
                else:
                    df = pd.DataFrame([vendedores])
            else:
                df = vendedores.copy()

            df['CONSORCIADO'] = consorciado
            df['DATA VENDA'] = data['data_venda']
            dfs.append(df)

        return pd.concat(dfs, ignore_index=True)