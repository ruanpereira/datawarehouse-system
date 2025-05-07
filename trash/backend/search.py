"""
Sistema de Geração de Relatórios de Vendas

Módulos:
- tkinter: Interface gráfica
- pandas: Manipulação de dados
- PyPDF2: Leitura de arquivos PDF
- python-docx: Geração de documentos Word
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import pandas as pd
import os
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Pt

class RelatorioVendas:
    """Classe responsável pela geração de relatórios de vendas em formato DOCX.
    
    Attributes:
        doc (Document): Objeto do documento Word
        dados (dict): Dicionário com os dados para preenchimento do relatório
    """
    
    def __init__(self, dados):
        """Inicializa o documento e configura os estilos base
        
        Args:
            dados (dict): Dados estruturados para o relatório
        """
        self.doc = Document()
        self.dados = dados
        self._configurar_estilos()

    def _configurar_estilos(self):
        """Configura os estilos padrão de fonte para o documento"""
        style = self.doc.styles['Normal']
        font = style.font  # type: ignore
        font.name = 'Arial'  # Define fonte Arial
        font.size = Pt(11)   # Define tamanho 11pt

    def add_title(self, text):
        """Adiciona um título principal ao documento
        
        Args:
            text (str): Texto do título
        """
        p = self.doc.add_heading(level=1)  # Cria cabeçalho nível 1 (título principal)
        run = p.add_run(text)              # Adiciona o texto
        run.bold = True                    # Aplica negrito

    def add_subtitle(self, text):
        """Adiciona um subtítulo ao documento
        
        Args:
            text (str): Texto do subtítulo
        """
        p = self.doc.add_heading(level=2)  # Cria cabeçalho nível 2 (subtítulo)
        run = p.add_run(text)
        run.bold = True

    def criar_relatorio(self):
        """Estrutura e gera todo o documento do relatório"""
        # Adiciona título principal
        self.add_title("Relatório de Vendas")
        
        # Adiciona informações do vendedor
        self.doc.add_paragraph(f"Vendedor: {self.dados['nome_vendedor']}")
        
        # Seção de resumo geral
        self.add_subtitle("Resumo Geral")
        self.doc.add_paragraph(f"Média total de vendas: {self.dados['media_total']}")
        
        # Tabelas de médias
        self._criar_tabela("Média por Região", 'media_regiao')
        self._criar_tabela("Média por Período", 'media_periodo')
        
        # Seções condicionais
        if self.dados.get('media_data'):
            self._criar_tabela("Média por Data", 'media_data')
        if self.dados.get('media_geral'):
            self.add_subtitle("Média Geral de Todos os Vendedores")
            self.doc.add_paragraph(f"Média geral: {self.dados['media_geral']}")
        
        # Observações finais
        self.add_subtitle("Observações Finais")
        self.doc.add_paragraph("[Análises adicionais e comentários]")
        
        # Salva o documento
        self.doc.save('relatorio_vendas_completo.docx')

    def _criar_tabela(self, titulo, chave):
        """Cria uma tabela formatada no documento
        
        Args:
            titulo (str): Título da seção da tabela
            chave (str): Chave do dicionário de dados para os valores da tabela
        """
        self.add_subtitle(titulo)
        
        # Cria tabela com 2 colunas
        table = self.doc.add_table(rows=1, cols=2)
        
        # Cabeçalho da tabela
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Categoria'
        hdr_cells[1].text = 'Valor Médio'
        
        # Preenche as linhas com os dados
        for chave_item, valor in self.dados[chave].items():
            row_cells = table.add_row().cells
            row_cells[0].text = chave_item
            row_cells[1].text = valor


class SearchingFile:
    """Classe principal para manipulação de dados e geração de relatórios
    
    Attributes:
        dataset (DataFrame): Dados carregados do arquivo
        columns_list (list): Lista de colunas disponíveis
        numeric_columns (list): Colunas numéricas identificadas
        relatorio_data (dict): Dados processados para o relatório
    """
    
    def __init__(self):
        """Inicializa o objeto com valores padrão"""
        self.dataset = None          # DataFrame com os dados carregados
        self.columns_list = []       # Lista de colunas disponíveis
        self.numeric_columns = []    # Colunas numéricas identificadas
        self.relatorio_data = {}     # Dados estruturados para o relatório

    def load_dataset(self, file_path):
        """Carrega dados de diferentes formatos de arquivo
        
        Args:
            file_path (str): Caminho completo do arquivo a ser carregado
            
        Returns:
            bool: True se o carregamento for bem sucedido, False caso contrário
            
        Raises:
            ValueError: Se o formato do arquivo não for suportado
        """
        # Extrai a extensão do arquivo
        ext = os.path.splitext(file_path)[1].lower()
        
        try:
            # Carrega o arquivo conforme seu formato
            if ext == '.csv':
                self.dataset = pd.read_csv(file_path)
            elif ext in ('.xlsx', '.xls'):
                self.dataset = pd.read_excel(file_path)
            elif ext == '.pdf':
                # Para PDFs, extrai todo o texto
                reader = PdfReader(file_path)
                text = "\n".join([page.extract_text() for page in reader.pages])
                self.dataset = pd.DataFrame({'Texto': [text]})
            else:
                raise ValueError("Formato não suportado")
            
            # Analisa as colunas após carregar
            self._analisar_colunas()
            return True
            
        except Exception as e:
            # Exibe mensagem de erro em caso de falha
            messagebox.showerror("Erro", f"Falha ao ler arquivo: {str(e)}")
            return False

    def _analisar_colunas(self):
        """Analisa e classifica as colunas do dataset
        1. Identifica todas as colunas
        2. Detecta colunas numéricas
        3. Tenta converter colunas com nomes comuns para numéricas
        """
        # Lista todas as colunas
        self.columns_list = self.dataset.columns.tolist()
        
        # Identifica colunas numéricas
        self.numeric_columns = self.dataset.select_dtypes(include=['number']).columns.tolist()
        
        # Se não encontrou colunas numéricas, tenta converter algumas comuns
        if not self.numeric_columns:
            for col in ['valor', 'venda', 'preço', 'total']:
                if col in [c.lower() for c in self.columns_list]:
                    # Encontra o nome real da coluna (case sensitive)
                    col_real = [c for c in self.columns_list if c.lower() == col][0]
                    try:
                        # Tenta converter para numérico
                        self.dataset[col_real] = pd.to_numeric(self.dataset[col_real])
                        self.numeric_columns.append(col_real)
                    except:
                        continue  # Ignora colunas que não podem ser convertidas

    def preparar_relatorio(self, coluna, valor, coluna_media=None, agrupamentos=None):
        """Prepara e processa os dados para geração do relatório
        
        Args:
            coluna (str): Nome da coluna para filtragem
            valor (str): Valor específico para filtrar na coluna
            coluna_media (str, optional): Coluna numérica para cálculos
            agrupamentos (list, optional): Lista de colunas para agrupamento
            
        Raises:
            Exception: Se ocorrer erro durante o processamento
        """
        # Caso especial para PDF (apenas texto)
        if 'Texto' in self.columns_list:
            messagebox.showwarning("Aviso", "PDF não suporta análise detalhada. Gerando relatório básico.")
            self.relatorio_data = {'conteudo_pdf': self.dataset['Texto'].iloc[0]}
            self._gerar_relatorio_pdf()
            return

        # Define padrões se não especificado
        coluna_media = coluna_media or (self.numeric_columns[0] if self.numeric_columns else None)
        agrupamentos = agrupamentos or [c for c in self.columns_list if c != coluna]
        
        # Valida se existe coluna numérica para cálculos
        if not coluna_media:
            messagebox.showerror("Erro", "Nenhuma coluna numérica encontrada para cálculo de médias")
            return

        # Filtra os dados conforme seleção
        self.column_name = coluna
        self.selected_value = valor
        self.filtered_rows = self.dataset[self.dataset[self.column_name] == self.selected_value]
        
        # Prepara estrutura de dados para o relatório
        self.relatorio_data = {
            'filtro': f"{coluna}: {valor}",
            'dados_gerais': self._get_dados_gerais(coluna_media),
            'agrupamentos': {}
        }

        # Calcula médias para cada coluna de agrupamento
        for agrupamento in agrupamentos:
            if agrupamento in self.columns_list and agrupamento != coluna:
                self.relatorio_data['agrupamentos'][agrupamento] = self._formatar_medias(agrupamento, coluna_media)
        
        # Gera o documento final
        self._gerar_relatorio_docx()

    def _get_dados_gerais(self, coluna_media):
        """Calcula métricas gerais para o relatório
        
        Args:
            coluna_media (str): Coluna numérica base para cálculos
            
        Returns:
            dict: Dicionário com métricas formatadas
        """
        return {
            'média': f"R$ {self.filtered_rows[coluna_media].mean():.2f}",
            'média_geral': f"R$ {self.dataset[coluna_media].mean():.2f}",
            'total': f"R$ {self.filtered_rows[coluna_media].sum():.2f}",
            'contagem': len(self.filtered_rows)
        }

    def _formatar_medias(self, agrupamento, coluna_media):
        """Calcula médias agrupadas e formata os resultados
        
        Args:
            agrupamento (str): Coluna para agrupamento dos dados
            coluna_media (str): Coluna numérica para cálculo das médias
            
        Returns:
            dict: Dicionário com os resultados formatados
        """
        try:
            return self.filtered_rows.groupby(agrupamento)[coluna_media].mean().apply(
                lambda x: f"R$ {x:.2f}" if pd.notnull(x) else "N/A"
            ).to_dict()
        except:
            return {"Erro": "Não foi possível calcular agrupamento"}

    def _gerar_relatorio_docx(self):
        """Gera o relatório detalhado em formato DOCX"""
        doc = Document()
        
        # Cabeçalho do documento
        doc.add_heading('Relatório Analítico', 0)
        doc.add_paragraph(f"Filtro aplicado: {self.relatorio_data['filtro']}")
        
        # Seção de métricas gerais
        doc.add_heading('Métricas Gerais', level=1)
        for k, v in self.relatorio_data['dados_gerais'].items():
            doc.add_paragraph(f"{k.replace('_', ' ').title()}: {v}")
        
        # Seções de agrupamento
        for agrupamento, dados in self.relatorio_data['agrupamentos'].items():
            doc.add_heading(f"Por {agrupamento}", level=2)
            
            # Cria tabela para cada agrupamento
            table = doc.add_table(rows=1, cols=2)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = agrupamento
            hdr_cells[1].text = 'Média'
            
            # Preenche as linhas da tabela
            for categoria, media in dados.items():
                row_cells = table.add_row().cells
                row_cells[0].text = str(categoria)
                row_cells[1].text = media
        
        # Salva o documento e informa o usuário
        doc.save('relatorio_analitico.docx')
        messagebox.showinfo("Sucesso", "Relatório gerado com sucesso!")

    def _gerar_relatorio_pdf(self):
        """Gera um relatório básico a partir de conteúdo de PDF"""
        doc = Document()
        doc.add_heading('Relatório Básico PDF', 0)
        
        # Limita o texto a 5000 caracteres para evitar problemas
        doc.add_paragraph(self.relatorio_data['conteudo_pdf'][:5000]) 
        
        # Salva e informa o usuário
        doc.save('relatorio_pdf_basico.docx')
        messagebox.showinfo("Sucesso", "Relatório básico do PDF gerado!")


class App(tk.Tk):
    """Classe principal da interface gráfica
    
    Attributes:
        searching (SearchingFile): Instância do processador de dados
    """
    
    def __init__(self):
        """Configura a janela principal e inicia a interface"""
        super().__init__()
        
        # Configurações básicas da janela
        self.title("Gerador de Relatório de Vendas")
        self.geometry("600x450")
        
        # Instancia o processador de dados
        self.searching = SearchingFile()
        
        # Cria os componentes da interface
        self._criar_widgets()

    def _criar_widgets(self):
        """Cria e posiciona todos os componentes da interface"""
        # Frame principal
        frame = ttk.Frame(self, padding=10)
        frame.pack(expand=True, fill='both')
        
        # Componentes para seleção de arquivo
        ttk.Label(frame, text="Selecione o arquivo de dados:").grid(row=0, column=0, pady=5, sticky='w')
        self.botao_arquivo = ttk.Button(frame, text="Procurar Arquivo", command=self.carregar_arquivo)
        self.botao_arquivo.grid(row=1, column=0, pady=5, sticky='ew')
        
        # Componentes para seleção de coluna e valor
        ttk.Label(frame, text="Coluna para filtrar:").grid(row=2, column=0, pady=5, sticky='w')
        self.combo_colunas = ttk.Combobox(frame, state="readonly")
        self.combo_colunas.grid(row=3, column=0, pady=5, sticky='ew')
        self.combo_colunas.bind("<<ComboboxSelected>>", self.carregar_valores)
        
        ttk.Label(frame, text="Valor para filtrar:").grid(row=4, column=0, pady=5, sticky='w')
        self.combo_valores = ttk.Combobox(frame, state="readonly")
        self.combo_valores.grid(row=5, column=0, pady=5, sticky='ew')
        
        # Opções adicionais
        self.chk_media_data = tk.BooleanVar()
        self.chk_media_geral = tk.BooleanVar()
        
        ttk.Checkbutton(frame, text="Incluir média por data", variable=self.chk_media_data).grid(
            row=6, column=0, pady=5, sticky='w')
        ttk.Checkbutton(frame, text="Incluir média geral", variable=self.chk_media_geral).grid(
            row=7, column=0, pady=5, sticky='w')
        
        # Botão principal
        ttk.Button(frame, text="Gerar Relatório", command=self.gerar_relatorio).grid(
            row=8, column=0, pady=20, sticky='ew')

    def carregar_arquivo(self):
        """Gerencia a seleção de arquivo pelo usuário"""
        # Tipos de arquivo suportados
        tipos_arquivo = [
            ("Arquivos Suportados", "*.csv *.xlsx *.xls *.pdf"),
            ("CSV", "*.csv"), 
            ("Excel", "*.xlsx *.xls"),
            ("PDF", "*.pdf")
        ]
        
        # Diálogo para seleção do arquivo
        file_path = filedialog.askopenfilename(filetypes=tipos_arquivo)
        
        # Se selecionou um arquivo válido, carrega os dados
        if file_path and self.searching.load_dataset(file_path):
            self.combo_colunas['values'] = self.searching.columns_list
            self._atualizar_status(f"Arquivo carregado: {os.path.basename(file_path)}")

    def carregar_valores(self, event):
        """Carrega os valores disponíveis para a coluna selecionada
        
        Args:
            event: Evento de seleção da combobox (não utilizado)
        """
        coluna = self.combo_colunas.get()
        
        # Se existe uma coluna selecionada e dados carregados
        if coluna and hasattr(self.searching, 'dataset'):
            # Obtém valores únicos da coluna, ordenados e sem valores nulos
            valores = sorted(self.searching.dataset[coluna].dropna().unique().tolist())
            self.combo_valores['values'] = valores
            self._atualizar_status(f"{len(valores)} valores encontrados na coluna {coluna}")

    def gerar_relatorio(self):
        """Coordena o processo de geração do relatório"""
        # Valida as seleções antes de prosseguir
        if not self._validar_selecoes():
            return
        
        try:
            # Solicita a geração do relatório
            self.searching.preparar_relatorio(
                self.combo_colunas.get(),
                self.combo_valores.get(),
                self.chk_media_data.get(),
                self.chk_media_geral.get()
            )
        except Exception as e:
            # Exibe mensagem em caso de erro
            messagebox.showerror("Erro", f"Falha ao gerar relatório: {str(e)}")

    def _validar_selecoes(self):
        """Valida as seleções do usuário antes de gerar o relatório
        
        Returns:
            bool: True se as seleções são válidas, False caso contrário
        """
        if not (self.combo_colunas.get() and self.combo_valores.get()):
            messagebox.showerror("Erro", "Selecione a coluna e o valor para filtrar!")
            return False
        return True

    def _atualizar_status(self, mensagem):
        """Atualiza o texto do botão com feedback do status
        
        Args:
            mensagem (str): Texto completo a ser exibido
        """
        # Limita o texto a 30 caracteres para caber no botão
        self.botao_arquivo['text'] = mensagem[:30] + "..." if len(mensagem) > 30 else mensagem


if __name__ == "__main__":
    # Ponto de entrada principal - cria e executa a aplicação
    app = App()
    app.mainloop()